import os, sys, numpy as np
from docx import Document

results_dir = r'D:\Projects\lpSurface\OCCT\results'
doc = Document()

# 设置默认字体为宋体
style = doc.styles['Normal']
style.font.name = '宋体'

doc.add_heading('NURBS 曲面直纹面分区报告', 0)
doc.add_paragraph(f'曲面类型：随机生成  |  平滑迭代次数：5')
doc.add_paragraph('')

# ── 零、NURBS 曲面参数 ──
rm = os.path.join(results_dir, 'run_meta.txt')
meta = {}
if os.path.exists(rm):
    with open(rm) as f:
        for line in f:
            if '=' in line:
                k, v = line.strip().split('=', 1)
                meta[k.strip()] = v.strip()

doc.add_heading('NURBS 曲面参数', 1)
if meta:
    m_table = doc.add_table(rows=0, cols=2, style='Table Grid')
    for k, v in meta.items():
        row = m_table.add_row().cells
        row[0].text = k
        row[1].text = v
    doc.add_paragraph('')
doc.add_paragraph(
    'NURBS 曲面定义：B 样条张量积曲面，控制点网格 '
    f'{meta.get("nurbs_ctrl_u","?")}×{meta.get("nurbs_ctrl_v","?")}，'
    f'次数 {meta.get("nurbs_degree_u","?")}×{meta.get("nurbs_degree_v","?")}，'
    f'参数域 U: [{meta.get("nurbs_domain_u","?.?")}]，'
    f'V: [{meta.get("nurbs_domain_v","?.?")}]。'
    '控制点 Z 坐标为 4~6 组正弦/余弦波叠加，幅度 0.08~0.35，频率 1.2~4.5，相位随机。'
    '网格采样 60×60 = 3721 顶点、7200 三角形。')

# ── 〇、容差定义 ──
doc.add_heading('容差定义', 1)
doc.add_paragraph(
    '直纹面 S(u,v) = (1-v)·γ₁(u) + v·γ₂(u)，其中 γ₁ 和 γ₂ 为分区边界环上'
    '两个分割点 P,Q 之间的两段弧（等弧长参数化，即 φ(u)=u）。\n\n'
    '优化目标：minimize Σ ||Q_k - S(u*_k, v*_k)||²，其中 (u*_k, v*_k) 为'
    '采样点 Q_k 在直纹面 60×24 离散网格上的最近投影点（包络定理：投影坐标'
    '无梯度，仅对 S 求导）。通过 PyTorch L-BFGS（strong Wolfe line search）'
    '优化分割点 t_P, t_Q 及弧长比例惩罚项（λ=50）。\n\n'
    '容差计算：在优化后的直纹面上密集采样（80×32 参数网格 = 2560 个顶点），'
    '对每个分区内部采样点（≤200 个面重心），计算其到直纹面网格的最近欧氏距离。\n'
    'maxDist = 所有内部点的最大距离\n'
    'rmsDist = sqrt(Σ d² / N)，均方根距离\n\n'
    '容差驱动迭代：若任一分区的 maxDist > tolTarget，则将平滑尺度 σ 缩小为 '
    '0.7 倍，重新执行拉普拉斯平滑（迭代次数 K = ceil((2σ/h)²)），再重新优化，'
    '最多 5 次重试。'
)
doc.add_heading('一、分区容差汇总', 1)
doc.add_paragraph('下表列出各分区经梯度优化后直纹面到分区内部采样点的最大和均方根距离。')

tol_file = os.path.join(results_dir, 'tolerance.txt')
if os.path.exists(tol_file):
    table = doc.add_table(rows=1, cols=5, style='Table Grid')
    headers = ['分区编号', 't_P', 't_Q', 'maxDist', 'rmsDist']
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    with open(tol_file) as f:
        f.readline()
        for line in f:
            parts = line.strip().split()
            if len(parts) >= 5:
                row = table.add_row().cells
                row[0].text = parts[0]
                row[1].text = parts[1][:6]
                row[2].text = parts[2][:6]
                row[3].text = f"{float(parts[3]):.4f}"
                row[4].text = f"{float(parts[4]):.4f}"
    doc.add_paragraph('')
    doc.add_paragraph(
        '参数说明：t_P 和 t_Q 为在分区边界多边形环上的分割点位置（取值范围 [0, 1]），'
        '直纹面由两段弧 γ₁(P→Q) 和 γ₂(Q→P) 通过等弧长参数化构建。'
        'maxDist 为直纹面四边形网格到分区内部采样点的最大欧氏距离，'
        'rmsDist 为均方根距离。')

# ── 二、各分区详情 ──
doc.add_heading('二、各分区详情', 1)
pid = 0
while True:
    lf = os.path.join(results_dir, f'part_{pid}_loop.txt')
    pf = os.path.join(results_dir, f'part_{pid}_points.txt')
    if not os.path.exists(lf):
        break
    loop = np.loadtxt(lf)
    pts = np.loadtxt(pf)
    if pts.ndim == 1:
        pts = pts.reshape(1, -1)

    doc.add_heading(f'分区 {pid}', 2)

    # 基本参数
    p_table = doc.add_table(rows=2, cols=3, style='Table Grid')
    p_table.rows[0].cells[0].text = '边界环顶点数'
    p_table.rows[0].cells[1].text = '内部采样点数'
    p_table.rows[0].cells[2].text = '环总弧长'
    p_table.rows[1].cells[0].text = str(loop.shape[0])
    p_table.rows[1].cells[1].text = str(pts.shape[0])
    segs = np.linalg.norm(np.diff(loop, axis=0), axis=1)
    p_table.rows[1].cells[2].text = f"{np.sum(segs):.3f}"
    doc.add_paragraph('')

    # 包围盒
    doc.add_paragraph(
        f'包围盒 X: [{loop[:,0].min():.3f}, {loop[:,0].max():.3f}]  '
        f'Y: [{loop[:,1].min():.3f}, {loop[:,1].max():.3f}]  '
        f'Z: [{loop[:,2].min():.3f}, {loop[:,2].max():.3f}]')

    # 分割点位置
    doc.add_paragraph(
        f'分割点 t_P: {loop.shape[0]} 个顶点环上约第 '
        f'{int(loop.shape[0] * (float(open(tol_file).readlines()[pid+1].split()[1]) if os.path.exists(tol_file) else 0.3))} 个顶点处')

    # 直纹面数据
    rf = os.path.join(results_dir, f'ruled_surf_{pid}.obj')
    if os.path.exists(rf):
        vc = fc = 0
        with open(rf) as fh:
            for line in fh:
                if line.startswith('v '): vc += 1
                elif line.startswith('f '): fc += 1
        doc.add_paragraph(f'直纹面 OBJ 文件：{vc} 个顶点（80×32 参数采样），{fc} 个四边形面')

    # 优化结果
    rf2 = os.path.join(results_dir, f'part_{pid}_ruled.txt')
    if os.path.exists(rf2):
        with open(rf2) as f:
            doc.add_paragraph('优化结果参数：')
            doc.add_paragraph(f.read().strip())
    pid += 1

# ── 三、平滑迭代历史 ──
doc.add_heading('三、拉普拉斯平滑迭代历史', 1)
doc.add_paragraph('每次迭代中边界顶点的最大位移和平均位移（UV 参数域单位）。')

sh = os.path.join(results_dir, 'smooth_history.txt')
if os.path.exists(sh):
    with open(sh) as f:
        lines = f.readlines()
    if len(lines) > 1:
        table = doc.add_table(rows=1, cols=3, style='Table Grid')
        for i, t in enumerate(['迭代编号', '最大位移 (maxDisp)', '平均位移 (avgDisp)']):
            table.rows[0].cells[i].text = t
        for line in lines[1:]:
            p = line.strip().split()
            if len(p) >= 3:
                row = table.add_row().cells
                row[0].text = p[0]
                row[1].text = f"{float(p[1]):.6f}"
                row[2].text = f"{float(p[2]):.6f}"

# ── 四、算法流程说明 ──
doc.add_heading('五、算法流程说明', 1)
doc.add_paragraph(
    '1. Hard-EM 分区算法将 NURBS 曲面网格顶点分配到 K=8 个分区，'
    '每个分区由一个直纹面近似。\n'
    '2. 极小区域合并：通过动态肘部阈值检测并合并孤立小分区。\n'
    '3. UV 参数域图拉普拉斯平滑：在参数域中迭代移动分区边界顶点，'
    '消除锯齿，保持外部顶点在域边界上。\n'
    '4. 调和网格更新：以平滑后边界为 Dirichlet 条件，求解 Δv=0，'
    '调整内部顶点使整体形变最小。\n'
    '5. 边界折线在分支点处分段提取。\n'
    '6. 直纹面优化：对每个分区边界环搜索最优分割点 t_P、t_Q，'
    '以等弧长参数化构建直纹面 S(u,v) = (1-v)γ₁(u) + vγ₂(u)，'
    '通过 PyTorch L-BFGS 最小化直纹面到内部点的投影距离。\n'
    '7. 容差计算：在直纹面上密集采样（80×32），计算各分区内部点'
    '到直纹面的最大及均方根距离。')

output = r'D:\Projects\lpSurface\OCCT\ruled_surface_report.docx'
# Fallback if file is locked
try:
    doc.save(output)
except PermissionError:
    import time
    output = r'D:\Projects\lpSurface\OCCT\ruled_surface_report_' + time.strftime('%H%M%S') + '.docx'
    doc.save(output)
print(f'中文报告已保存: {output}')
